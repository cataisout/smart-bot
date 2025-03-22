import fitz  # PyMuPDF
from docx import Document
from langchain.schema import Document as LangchainDocument
import io


def extract_text_from_pdf(file_bytes: io.BytesIO) -> str:
    """Extrai texto de um arquivo PDF carregado via BytesIO."""
    try:
        doc = fitz.open(stream=file_bytes, filetype="pdf")
        text = "\n".join(page.get_text("text") for page in doc)
        return text
    except Exception as e:
        return f"Erro ao extrair texto de PDF: {str(e)}"

def extract_text_from_docx(file_bytes: io.BytesIO) -> str:
    """Extrai texto de um arquivo DOCX carregado via BytesIO."""
    try:
        doc = docx.Document(file_bytes)
        text = "\n".join([para.text for para in doc.paragraphs])
        return text
    except Exception as e:
        return f"Erro ao extrair texto de DOCX: {str(e)}"

def extract_text_from_txt(file_bytes: io.BytesIO) -> str:
    """Extrai texto de um arquivo TXT carregado via BytesIO."""
    try:
        return file_bytes.getvalue().decode("utf-8")  # Lê os bytes e converte para string
    except Exception as e:
        return f"Erro ao extrair texto de TXT: {str(e)}"

def extract_text_from_document(uploaded_file) -> LangchainDocument:
    """
    Extrai texto de um arquivo carregado via Streamlit.
    :param uploaded_file: Objeto `BytesIO` vindo do `st.file_uploader()`.
    :return: Documento Langchain com o texto extraído.
    """

    file_extension = uploaded_file.name.lower().split(".")[-1]  # Obtém a extensão

    try:
        if file_extension == "pdf":
            text = extract_text_from_pdf(io.BytesIO(uploaded_file.read()))
        elif file_extension == "docx":
            text = extract_text_from_docx(io.BytesIO(uploaded_file.read()))
        elif file_extension == "txt":
            text = extract_text_from_txt(io.BytesIO(uploaded_file.read()))
        else:
            return LangchainDocument(page_content="Formato de arquivo não suportado.", metadata={})

        return LangchainDocument(page_content=text, metadata={"source": uploaded_file.name})

    except Exception as e:
        return LangchainDocument(page_content=f"Erro ao processar o documento: {str(e)}", metadata={})